from docx import Document

# Input and output file paths
input_file = "Python/data.txt"  # Replace with the path to your file
output_file = "Python/output2.docx"  # Output Word document

try:
    # Read the input file
    with open(input_file, "r") as file:
        lines = file.readlines()

    # Create a Word document
    document = Document()
    document.add_heading("Command and Description Table", level=1)

    # Add a table to the document
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Add table headers
    header_cells = table.rows[0].cells
    header_cells[0].text = "Command/Description"
    header_cells[1].text = "Details"

    # Process each line and add to the table
    for line in lines:
        line = line.strip()  # Remove leading/trailing whitespace

        if line == "":  # Skip empty lines
            continue

        if ":" in line:
            # Split lines containing a colon into two columns
            command, details = line.split(":", 1)
            row_cells = table.add_row().cells
            row_cells[0].text = command.strip()
            row_cells[1].text = details.strip()
        else:
            # Add lines without a colon to the first column only
            row_cells = table.add_row().cells
            row_cells[0].text = line
            row_cells[1].text = ""  # Leave the second column empty

    # Save the Word document
    document.save(output_file)
    print(f"Data successfully written to {output_file}")

except FileNotFoundError:
    print("The file data.txt does not exist.")
except Exception as e:
    print(f"An error occurred: {e}")
