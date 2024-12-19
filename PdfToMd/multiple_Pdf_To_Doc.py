import os
import PyPDF2
from docx import Document


def convert_pdf_to_doc(pdf_path, doc_path):
    # Step 1: Open the PDF file
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Step 2: Create a new Word document
        doc = Document()

        # Step 3: Extract text from each page and add to the Word document
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
                doc.add_paragraph("\n")  # Add spacing between pages

        # Step 4: Save the Word document
        doc.save(doc_path)

    print(f"Word document created: {doc_path}")


def convert_multiple_pdfs(pdf_dir, output_dir):
    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Loop through all PDF files in the input directory
    for file_name in os.listdir(pdf_dir):
        if file_name.endswith(".pdf"):
            pdf_path = os.path.join(pdf_dir, file_name)
            doc_path = os.path.join(output_dir, file_name.replace(".pdf", ".docx"))

            # Convert PDF to Word
            convert_pdf_to_doc(pdf_path, doc_path)


# Input and output directories
pdf_directory = "/home/mss/Desktop/ProblemSolving/Python/PdfToMd/Commands/"
output_directory = "/home/mss/Desktop/ProblemSolving/Python/PdfToMd/DocxFiles"

convert_multiple_pdfs(pdf_directory, output_directory)
