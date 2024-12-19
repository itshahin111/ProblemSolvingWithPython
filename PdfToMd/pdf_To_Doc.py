import PyPDF2
from docx import Document


def convert_pdf_to_doc(pdf_path, doc_path):
    # Step 1: Open the PDF file
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Step 2: Create a new Word document
        doc = Document()

        # Step 3: Extract text from each page and add to the Word document
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
                doc.add_paragraph("\n")  # Add spacing between pages

        # Step 4: Save the Word document
        doc.save(doc_path)

    print(f"Word document created: {doc_path}")


# Input PDF file path and output Word file path
pdf_file_path = "/home/mss/Desktop/ProblemSolving/Python/PdfToMd/Commands/LinuxBasicCommands.pdf"
word_file_path = "/home/mss/Desktop/ProblemSolving/Python/PdfToMd/Docx/LinuxBasicCommands.docx"

convert_pdf_to_doc(pdf_file_path, word_file_path)