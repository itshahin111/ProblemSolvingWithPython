from docx import Document

# Path to the input .odd file
# file_path = "LinuxCommands.txt"  # Replace with the path to your .odd file
output_docx = "output.docx"  # Path for the output Word document

try:
    # Read the .odd file
    with open(file_path, "r") as file:
        lines = file.readlines()

    # Process each line to split into command and description
    table_data = [line.strip().split(": ", 1) for line in lines if line.strip()]

    # Create a Word document
    document = Document()
    document.add_heading("Command and Description Table", level=1)

    # Add a table to the document
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Add table headers
    header_cells = table.rows[0].cells
    header_cells[0].text = "Command"
    header_cells[1].text = "Description"

    # Add rows to the table
    for command, description in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = command.strip()
        row_cells[1].text = description.strip()

    # Save the document
    document.save(output_docx)
    print(f"Word document successfully written to {output_docx}")

except FileNotFoundError:
    print("The file does not exist.")
except Exception as e:
    print(f"An error occurred: {e}")
